from flask import Flask, request, jsonify, send_file, session
import os
import tempfile
from docx.shared import Pt, Inches
from openai import Client
from docx import Document
from docx.enum.section import WD_SECTION
from dotenv import load_dotenv
from python_docx_replace import docx_replace
import re
from io import BytesIO
from docx.shared import Pt
from docx.oxml.ns import qn
import docx2txt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

app = Flask(__name__)
app.config['UPLOAD FOLDER'] = 'uploads'

load_dotenv()
key = os.getenv("OPENAI_API_KEY")



def set_font_style(document, font_name='Arial'):
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name) 

    
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = font_name
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def replace_placeholder(paragraph, placeholder, replacement):
    if paragraph.text == placeholder:
        combined_text = "".join([run.text for run in paragraph.runs])
        new_text = combined_text.replace(placeholder, replacement)
            
        # Clear existing runs and add the replaced text as a single run
        for run in paragraph.runs:
            run.text = ""
        paragraph.runs[0].text = new_text
    

def bold_text(paragraph, text):
    run = paragraph.add_run(text)
    run.bold = True


def gpt_response(old_doc_text, job_description):
    client = Client(api_key=key)

    prompt = f"""Here is the CV/Resume of a person, I want you to efficiently convert the following profile into the new format provided below. While converting make sure there are no "I", "He", "She" etc. change first-person and third-person tones to emphasize action words more. 
    I also need you to compare the following job decription (which is basically a requirement or need for a person) with the person's profile and give me a percentage of how likely the person
    is to be a match for the job. I also want you to give me the missing keywords that are present in the job description but not in the person's profile  
    DO NOT output anything other than the new format!
    Ensure the output is consistently structured as follows:

    Job Description:
    {job_description}

    Old format:
    {old_doc_text}

    New format (you ONLY need to output this):

    Resource Name: (Extract and display the profile's Resource Name as the heading)
    Percentage Match: (Percentage of how much the person matches the job)
    Missing Keywords: (Keywords that are present in job description but not in users profile)
    Job Role: (Extract and display the profile's Job Role from the summary as the heading)
    Summary: (Keep the summary the same)
    Key Skills & Competencies: (Identify and extract heading like "1. Technical Expertise" or "2. Technical Toolkits" from the profile and then elaborate them in 2-3 words after a colon (Start directly from action words, do not use the persons name in this section). KEEP THE HEADINGS GENERIC AND FOLLOW THE EXAMPLE BELOW
    Academic Summary: (Display as provided) - must be in bullet points. Ensure that "Masters" degrees always come before "Bachelors"
    Work Summary: (Provide in detailed paragraphs without altering any information. For each project, include the following details)
    - Project x: (where x is a number, followed by name of the project, this CANNOT be empty)
    - Environment: (Always identify and display the job role from the respective project separated by commas, this CANNOT be empty)
    - Outline: (For each project, give the overview of the project, this CANNOT be empty )
    - Responsibilities: (For each project, detail the responsibilities separated by commas, do NOT use points, this CANNOT be empty)


    Here is an example of Work Summary (it should strictly follow this format!):
    Project 3: ERP Sales – Desktop Application and Backend Services
    Environment: Visual Studio, Vb.Net, MS-SQL
    Outline: A desktop application based on ERP solution for enabling customers to book and track orders, manage deliveries, as well as generate invoices, schemes and pricing.
    Responsibilities: Developed and maintain desktop-based business applications built on ASP.NET. Worked on scripting and performed unit testing. Worked on WinForms and review integrations between applications. Implemented procedures for gathering and analysis of user and business needs for solutions. Worked on troubleshooting, debugging and resolution of production issues along with support services.
    

    Here is an example of Key Skills & Competencies (it should strictly follow this format!):
    1. Technical Toolkit: Swift, iOS, Flutter, XCode
    2. Business User Requirement Gathering: 6+ years of experience eliciting and documenting user requirements, translating them into clear and actionable BRDs. Proven ability to break down complex projects into manageable deliverables.
    3. Communication and Collaboration: Strong communication skills, evidenced by effective collaboration with clients and teams. Acted as a bridge between commercial and technical divisions

        """



    response = client.chat.completions.create(
        model='gpt-3.5-turbo-16k',
        messages=[{"role": "assistant", "content": [{"type": "text", "text": prompt}]}],
        max_tokens=2000,
        temperature=0.1,
        top_p=1.0
    )

    return response.choices[0].message.content





@app.route('/upload', methods=['POST'])
def process_file():
    if 'file' not in request.files or 'job_description' not in request.form:
        return jsonify({"error": "File and job description are required"}), 400

    file = request.files['file']
    job_description = request.form['job_description']

    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400

    file_stream = BytesIO(file.read())
    
    #old_doc = Document(file_stream)

    #Changing DOcument from docx lib to docx2txt lib
    old_doc_text = docx2txt.process(file_stream)
    
    print(old_doc_text)


    original_filename = os.path.splitext(file.filename)[0]
    new_filename = f"{original_filename}_CONVERTED.docx"
    new_filename = new_filename.replace(' ','')

    #old_doc.save("puranaDocument.docx")

    #old_doc_text = ""

    #for paragraph in old_doc.paragraphs:
    #    old_doc_text += paragraph.text + "\n"


    resp = gpt_response(old_doc_text, job_description)

    resp = resp.split("\n")

    if '' in resp:
        resp.remove('')



    text = "\n".join(resp)

    # Use regex to extract the required sections
    resource_name = re.search(r"Resource Name:\s*(.*)", text).group(1)
    job_role = re.search(r"Job Role:\s*(.*)", text).group(1)
    percentage_match = re.search(r"Percentage Match:\s*(.*)", text).group(1)
    missing_keywords = re.search(r"Missing Keywords:\s*(.*)", text).group(1)
    summary = re.search(r"Summary:\s*(.*)", text).group(1)
    key_skills_competencies = re.search(r"Key Skills & Competencies:\s*(.*?)Academic Summary:", text, re.DOTALL).group(1).strip()
    academic_summary = re.search(r"Academic Summary:\s*(.*?)Work Summary:", text, re.DOTALL).group(1).strip()
    work_summary = re.search(r"Work Summary:\s*(.*)", text, re.DOTALL).group(1).strip()


    print(work_summary + "\n\n\n")

    # Store in appropriate data structures
    profile_data = {
        "Resource Name": resource_name,
        "Job Role": job_role,
        "Percentage Match": percentage_match,
        "Missing Keywords": missing_keywords,
        "Summary": summary,
        "Key Skills & Competencies": key_skills_competencies.split("\n"),
        "Academic Summary": academic_summary.split("\n"),
        "Work Summary": []
    }

    profile_data["Academic Summary"] = [line.replace('-', '•') for line in profile_data["Academic Summary"]]

    profile_data["Academic Summary"].sort(key=lambda x: ("Bachelors" in x, "Masters" not in x))



    projects = re.findall(
    r"Project \d+: (.*?)\n\s*Environment:\s*(.*?)\n\s*Outline:\s*(.*?)\n\s*Responsibilities:\s*(.*?)(?=\nProject \d+:|$)",
    work_summary, re.DOTALL
    )

    # Populate profile_data with extracted projects
    for project in projects:
        responsibilities = project[3].strip().replace("- ", "").replace("\n", ". ")

        project_data = {
            "Project": project[0],
            "Environment": project[1],
            "Outline": project[2].strip(),
            "Responsibilities": responsibilities
        }
        profile_data["Work Summary"].append(project_data)


    for key, value in profile_data.items():
        print(f"{key}: {value}")

    
    new_file_path = os.path.join('template', 'Sample Profile 2.docx')
    new_file = Document(new_file_path)


    for paragraph in new_file.paragraphs:
        replace_placeholder(paragraph, "${resourcename}", profile_data["Resource Name"])

    for paragraph in new_file.paragraphs:    
        replace_placeholder(paragraph, "${jobrole}", profile_data["Job Role"])
        
    for paragraph in new_file.paragraphs:    
        replace_placeholder(paragraph, "${summary}", profile_data["Summary"])
        
    for paragraph in new_file.paragraphs:    
        replace_placeholder(paragraph, "${academic}", "\n".join(profile_data["Academic Summary"]))


    for paragraph in new_file.paragraphs:
        if "${keyskills}" in paragraph.text:
            paragraph.clear()
            for skill in profile_data["Key Skills & Competencies"]:
                skill_parts = skill.split(":")
                
                if len(skill_parts) == 2:
                    bold_text(paragraph, skill_parts[0] + ": ")
                    competencies = skill_parts[1].split(", ")
                    for competency in competencies:
                        competency = competency.strip()
                        paragraph.add_run("\n   • ")
                        bullet_point = paragraph.add_run(competency)
                        bullet_point.font.size = Pt(11)
                        
                    paragraph.add_run("\n")
                    paragraph.add_run("\n")  # Add new line after each skill section
        


    for paragraph in new_file.paragraphs:
        if "${worksummary}" in paragraph.text:
            paragraph.clear()
            for project in profile_data["Work Summary"]:
                table = new_file.add_table(rows=0, cols=2)
                
                # Add row for Environment and Project
                row_cells = table.add_row().cells
                
                # Environment column
                env_paragraph = row_cells[0].add_paragraph()
                env_run = env_paragraph.add_run("Environment")
                env_run.bold = True
                env_paragraph.add_run("\n" + project["Environment"])
                env_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                
                # Project column
                #proj_paragraph = row_cells[1].add_paragraph()
                #proj_run = proj_paragraph.add_run(project['Project'])
                #proj_run.bold = True
                #proj_paragraph.add_run("\n" + project["Outline"])
                #env_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                proj_paragraph = row_cells[1].add_paragraph()
                proj_run = proj_paragraph.add_run(project['Project'])  # Add project title
                proj_run.bold = True
                proj_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Unjustified
                
                outline_paragraph = row_cells[1].add_paragraph()  # New paragraph for outline
                outline_paragraph.add_run(project["Outline"])  # Add outline text
                outline_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # Justify outline
    


                # Add row for Responsibilities
                row_cells = table.add_row().cells
                row_cells[0].width = Inches(4)  # Adjust env column width
                row_cells[1].width = Inches(7)  # Adjust column width
                #row_cells[1].text = "Responsibilities:" "\n" + project["Responsibilities"]
                
                responsibilities_paragraph = row_cells[1].add_paragraph()
                responsibilities_run = responsibilities_paragraph.add_run("Responsibilities: ")
                responsibilities_run.bold = True
                responsibilities_paragraph.add_run("\n" + project["Responsibilities"])
                responsibilities_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY


                # Add space and border between tables
                new_file.add_paragraph("\n")




    set_font_style(new_file, font_name='Arial')

    #new_file.save('NewSampleFormatCV.docx')
    
    #print(profile_data["Work Summary"])
    #print(profile_data["Missing Keywords"])

    #Save new file to a temporary location
    temp_dir = tempfile.gettempdir()
    new_file_path = os.path.join(temp_dir, new_filename)
    new_file.save(new_file_path)


    #releasing old_doc_text buffer
    old_doc_text=''

    # Return the match percentage, missing keywords, and download link
    return jsonify({
        "percentage_match": profile_data["Percentage Match"],
        "missing_keywords": profile_data["Missing Keywords"],
        "download_link": f"/download?file_path={new_file_path}"
    })


@app.route('/download', methods=['GET'])
def download_file():
    filename = request.args.get('new_filename')
    file_path = request.args.get('file_path')
    return send_file(file_path, as_attachment=True, download_name=filename)

#testing
@app.route('/hello', methods=['GET'])
def hello_world():
    return "hello world"



if __name__ == '__main__':
    app.run(debug=True)